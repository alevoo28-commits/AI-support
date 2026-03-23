"""
knowledge_base.py
=================
Base de conocimiento por áreas de empresa.
Permite subir documentos (PDF, DOCX, XLSX, TXT), los parsea, los trocea
y construye un índice FAISS por área para búsqueda semántica.
Si no hay embeddings disponibles, usa búsqueda por palabras clave (BM25-like).
"""

from __future__ import annotations

import json
import os
import re
import time
import uuid
import datetime as _dt
from pathlib import Path
from typing import Optional

# ── Carpeta raíz de almacenamiento ──────────────────────────────────────────
_DEFAULT_BASE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "knowledge_base",
)

# ── Helpers de parseo ────────────────────────────────────────────────────────

def _parse_pdf(path: str) -> str:
    try:
        import pdfplumber
        text_parts: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n\n".join(text_parts)
    except Exception:
        pass
    try:
        import PyPDF2
        text_parts = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n\n".join(text_parts)
    except Exception as e:
        raise RuntimeError(f"No se pudo parsear PDF: {e}")


def _parse_docx(path: str) -> str:
    try:
        import docx
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # También extraer tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        raise RuntimeError(f"No se pudo parsear DOCX: {e}")


def _parse_excel(path: str) -> str:
    try:
        import pandas as pd
        xl = pd.ExcelFile(path)
        parts: list[str] = []
        for sheet in xl.sheet_names:
            df = xl.parse(sheet).fillna("")
            parts.append(f"=== Hoja: {sheet} ===")
            parts.append(df.to_string(index=False))
        return "\n\n".join(parts)
    except Exception as e:
        raise RuntimeError(f"No se pudo parsear Excel: {e}")


def _parse_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _parse_file(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(path)
    elif ext in (".docx", ".doc"):
        return _parse_docx(path)
    elif ext in (".xlsx", ".xls", ".xlsm"):
        return _parse_excel(path)
    elif ext in (".txt", ".md", ".csv"):
        return _parse_txt(path)
    else:
        raise ValueError(f"Tipo de archivo no soportado: {ext}")


def _split_text(text: str, chunk_size: int = 800, overlap: int = 150) -> list[str]:
    """Trocea texto en chunks con solapamiento."""
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size, chunk_overlap=overlap, separators=["\n\n", "\n", ". ", " ", ""]
        )
        return splitter.split_text(text)
    except Exception:
        # Fallback manual
        words = text.split()
        chunks: list[str] = []
        step = chunk_size - overlap
        for i in range(0, max(1, len(words) - overlap), step):
            chunk = " ".join(words[i : i + chunk_size])
            if chunk.strip():
                chunks.append(chunk)
        return chunks if chunks else [text[:4000]]


# ── Clase principal ──────────────────────────────────────────────────────────

class KnowledgeBaseManager:
    """Gestiona la base de conocimiento organizada por áreas."""

    METADATA_FILE = "metadata.json"
    DOCS_DIR = "documents"
    INDEX_DIR = "index"
    CHUNKS_FILE = "chunks.json"

    def __init__(self, base_dir: Optional[str] = None):
        self.base_dir = Path(base_dir or _DEFAULT_BASE)
        self.base_dir.mkdir(parents=True, exist_ok=True)
        self._areas_meta_path = self.base_dir / "areas.json"
        if not self._areas_meta_path.exists():
            self._write_areas({})
        self._bootstrap_legacy_areas_if_needed()

    def _bootstrap_legacy_areas_if_needed(self) -> None:
        """Reconstruye areas.json cuando hay datos legados en carpetas por área.

        Escenario típico: existen carpetas en knowledge_base/* con metadata/chunks,
        pero areas.json está vacío, por lo que la UI no muestra áreas.
        """
        areas = self._read_areas()
        if areas:
            return

        discovered: dict[str, dict] = {}
        for child in self.base_dir.iterdir():
            if not child.is_dir():
                continue

            area_id = child.name
            meta_path = child / self.METADATA_FILE
            name = area_id
            description = ""
            created_at = time.time()

            if meta_path.exists():
                try:
                    payload = json.loads(meta_path.read_text(encoding="utf-8"))
                    if isinstance(payload, dict):
                        name = str(payload.get("name") or name).strip() or name
                        description = str(payload.get("description") or "").strip()
                        created_raw = payload.get("created")
                        if isinstance(created_raw, str) and created_raw.strip():
                            try:
                                dt = _dt.datetime.strptime(created_raw.strip(), "%Y-%m-%d")
                                created_at = dt.timestamp()
                            except Exception:
                                pass
                except Exception:
                    pass

            # Solo agregar áreas con señales de contenido en formato actual o legado.
            has_current = (child / self.DOCS_DIR).exists() or (child / self.INDEX_DIR).exists()
            has_legacy = (child / "chunks.json").exists() or meta_path.exists()
            if not (has_current or has_legacy):
                continue

            discovered[area_id] = {
                "id": area_id,
                "name": name,
                "description": description,
                "created_at": created_at,
            }

        if discovered:
            self._write_areas(discovered)

    # ── Áreas ────────────────────────────────────────────────────────────────

    def _read_areas(self) -> dict:
        try:
            return json.loads(self._areas_meta_path.read_text(encoding="utf-8"))
        except Exception:
            return {}

    def _write_areas(self, data: dict) -> None:
        self._areas_meta_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    def _build_children_map(self, areas: dict) -> dict[str, list[str]]:
        children: dict[str, list[str]] = {}
        for area_id, meta in areas.items():
            parent = str(meta.get("parent_id") or "").strip()
            if not parent:
                continue
            children.setdefault(parent, []).append(area_id)
        for parent_id in children:
            children[parent_id].sort()
        return children

    def _area_depth(self, areas: dict, area_id: str) -> int:
        depth = 0
        current = area_id
        seen: set[str] = set()
        while True:
            if current in seen:
                break
            seen.add(current)
            meta = areas.get(current) or {}
            parent = str(meta.get("parent_id") or "").strip()
            if not parent or parent not in areas:
                break
            depth += 1
            current = parent
        return depth

    def _area_full_path(self, areas: dict, area_id: str) -> str:
        parts: list[str] = []
        current = area_id
        seen: set[str] = set()
        while True:
            if current in seen:
                break
            seen.add(current)
            meta = areas.get(current) or {}
            name = str(meta.get("name") or current).strip() or current
            parts.append(name)
            parent = str(meta.get("parent_id") or "").strip()
            if not parent or parent not in areas:
                break
            current = parent
        parts.reverse()
        return " > ".join(parts)

    def list_areas(self) -> list[dict]:
        areas = self._read_areas()
        result = []
        for area_id, meta in areas.items():
            docs = self.list_documents(area_id)
            depth = self._area_depth(areas, area_id)
            full_path = self._area_full_path(areas, area_id)
            result.append({
                **meta,
                "id": area_id,
                "doc_count": len(docs),
                "depth": depth,
                "full_path": full_path,
            })
        return sorted(result, key=lambda x: (x.get("full_path", ""), x.get("name", "")))

    def create_area(self, name: str, description: str = "", parent_id: Optional[str] = None) -> dict:
        name = name.strip()
        if not name:
            raise ValueError("El nombre del área no puede estar vacío.")
        area_id = re.sub(r"[^a-z0-9_]", "_", name.lower())[:48] + "_" + str(int(time.time()))[-5:]
        areas = self._read_areas()
        parent_clean = str(parent_id or "").strip() or None
        if parent_clean and parent_clean not in areas:
            raise ValueError("El área padre seleccionada no existe.")
        # Evitar duplicados por nombre
        for meta in areas.values():
            if meta.get("name", "").lower() == name.lower():
                raise ValueError(f"Ya existe un área con el nombre '{name}'.")
        areas[area_id] = {
            "id": area_id,
            "name": name,
            "description": description,
            "created_at": time.time(),
            "parent_id": parent_clean,
        }
        self._write_areas(areas)
        # Crear estructura de directorios
        area_path = self.base_dir / area_id
        (area_path / self.DOCS_DIR).mkdir(parents=True, exist_ok=True)
        (area_path / self.INDEX_DIR).mkdir(parents=True, exist_ok=True)
        return areas[area_id]

    def update_area_parent(self, area_id: str, parent_id: Optional[str]) -> dict:
        """Mueve un área dentro del árbol (re-parent)."""
        areas = self._read_areas()
        if area_id not in areas:
            raise ValueError("El área seleccionada no existe.")

        parent_clean = str(parent_id or "").strip() or None
        if parent_clean == area_id:
            raise ValueError("Un área no puede ser padre de sí misma.")
        if parent_clean and parent_clean not in areas:
            raise ValueError("El área padre seleccionada no existe.")

        # Evitar ciclos: el nuevo padre no puede ser descendiente del área.
        if parent_clean:
            descendants = self.get_related_area_ids([area_id], include_ancestors=False, include_descendants=True)
            descendants.discard(area_id)
            if parent_clean in descendants:
                raise ValueError("Movimiento inválido: el área padre es descendiente del área actual.")

        areas[area_id]["parent_id"] = parent_clean
        self._write_areas(areas)
        return areas[area_id]

    def get_related_area_ids(
        self,
        seed_area_ids: list[str],
        *,
        include_ancestors: bool = True,
        include_descendants: bool = True,
    ) -> set[str]:
        areas = self._read_areas()
        children = self._build_children_map(areas)

        related: set[str] = set(a for a in seed_area_ids if a in areas)

        if include_ancestors:
            for area_id in list(related):
                current = area_id
                seen: set[str] = set()
                while True:
                    if current in seen:
                        break
                    seen.add(current)
                    parent = str((areas.get(current) or {}).get("parent_id") or "").strip()
                    if not parent or parent not in areas:
                        break
                    related.add(parent)
                    current = parent

        if include_descendants:
            queue = list(related)
            seen_desc: set[str] = set(queue)
            while queue:
                current = queue.pop(0)
                for child in children.get(current, []):
                    if child in seen_desc:
                        continue
                    seen_desc.add(child)
                    related.add(child)
                    queue.append(child)

        return related

    def delete_area(self, area_id: str) -> bool:
        areas = self._read_areas()
        if area_id not in areas:
            return False

        children = self._build_children_map(areas)
        if children.get(area_id):
            raise ValueError("No puedes eliminar un área que tiene subáreas.")

        areas.pop(area_id)
        self._write_areas(areas)
        import shutil
        area_path = self.base_dir / area_id
        if area_path.exists():
            shutil.rmtree(area_path)
        return True

    def get_area(self, area_id: str) -> Optional[dict]:
        return self._read_areas().get(area_id)

    # ── Documentos ───────────────────────────────────────────────────────────

    def _doc_meta_path(self, area_id: str) -> Path:
        return self.base_dir / area_id / self.METADATA_FILE

    def list_documents(self, area_id: str) -> list[dict]:
        path = self._doc_meta_path(area_id)
        if path.exists():
            try:
                data = json.loads(path.read_text(encoding="utf-8"))
                if isinstance(data, list):
                    return data
                # Compatibilidad legado: metadata.json con objeto de área (no lista de docs)
                # Sintetizamos una lista de documentos desde chunks.json si existe.
                if isinstance(data, dict):
                    legacy_chunks_path = self.base_dir / area_id / "chunks.json"
                    if legacy_chunks_path.exists():
                        try:
                            legacy = json.loads(legacy_chunks_path.read_text(encoding="utf-8"))
                            if isinstance(legacy, dict):
                                filename = str(legacy.get("filename") or "documento_legacy").strip() or "documento_legacy"
                                chunks = legacy.get("chunks") or []
                                total_chunks = int(legacy.get("total_chunks") or (len(chunks) if isinstance(chunks, list) else 0))
                                return [
                                    {
                                        "id": "legacy_doc",
                                        "filename": filename,
                                        "stored_name": filename,
                                        "uploaded_at": float((self.get_area(area_id) or {}).get("created_at") or time.time()),
                                        "size_bytes": 0,
                                        "chunk_count": total_chunks,
                                    }
                                ]
                        except Exception:
                            pass
            except Exception:
                pass
        return []

    def _save_doc_meta(self, area_id: str, docs: list[dict]) -> None:
        path = self._doc_meta_path(area_id)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(json.dumps(docs, ensure_ascii=False, indent=2), encoding="utf-8")

    def upload_document(self, area_id: str, filename: str, file_bytes: bytes) -> dict:
        """Guarda el archivo, lo parsea, trocea e indexa."""
        area_path = self.base_dir / area_id
        docs_path = area_path / self.DOCS_DIR
        docs_path.mkdir(parents=True, exist_ok=True)

        # Guardar archivo
        doc_id = str(uuid.uuid4())[:12]
        safe_name = re.sub(r"[^a-zA-Z0-9._\-]", "_", filename)
        stored_name = f"{doc_id}_{safe_name}"
        file_path = docs_path / stored_name
        file_path.write_bytes(file_bytes)

        # Parsear texto
        try:
            text = _parse_file(str(file_path))
        except Exception as e:
            file_path.unlink(missing_ok=True)
            raise RuntimeError(f"Error al parsear '{filename}': {e}")

        # Trocear y guardar chunks
        chunks = _split_text(text)
        chunks_path = area_path / self.INDEX_DIR / f"{doc_id}_chunks.json"
        chunks_path.parent.mkdir(parents=True, exist_ok=True)
        chunks_path.write_text(
            json.dumps(
                [{"doc_id": doc_id, "filename": filename, "chunk_idx": i, "text": c} for i, c in enumerate(chunks)],
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        # Metadatos
        doc_meta = {
            "id": doc_id,
            "filename": filename,
            "stored_name": stored_name,
            "uploaded_at": time.time(),
            "size_bytes": len(file_bytes),
            "chunk_count": len(chunks),
        }
        docs = self.list_documents(area_id)
        docs.append(doc_meta)
        self._save_doc_meta(area_id, docs)

        # Invalidar índice FAISS (se reconstruirá en próxima búsqueda)
        self._invalidate_faiss(area_id)

        return doc_meta

    def delete_document(self, area_id: str, doc_id: str) -> bool:
        docs = self.list_documents(area_id)
        new_docs = [d for d in docs if d["id"] != doc_id]
        if len(new_docs) == len(docs):
            return False
        removed = next(d for d in docs if d["id"] == doc_id)
        self._save_doc_meta(area_id, new_docs)

        # Borrar archivo y chunks
        file_path = self.base_dir / area_id / self.DOCS_DIR / removed.get("stored_name", "")
        if file_path.exists():
            file_path.unlink()
        chunks_path = self.base_dir / area_id / self.INDEX_DIR / f"{doc_id}_chunks.json"
        if chunks_path.exists():
            chunks_path.unlink()
        self._invalidate_faiss(area_id)
        return True

    # ── Chunks ───────────────────────────────────────────────────────────────

    def _load_all_chunks(self, area_id: str) -> list[dict]:
        index_path = self.base_dir / area_id / self.INDEX_DIR
        chunks: list[dict] = []

        # Formato actual
        if index_path.exists():
            for f in index_path.glob("*_chunks.json"):
                try:
                    payload = json.loads(f.read_text(encoding="utf-8"))
                    if isinstance(payload, list):
                        chunks.extend(payload)
                except Exception:
                    pass

        # Formato legado: area/chunks.json con {filename, chunks:[...]}
        if not chunks:
            legacy_path = self.base_dir / area_id / "chunks.json"
            if legacy_path.exists():
                try:
                    payload = json.loads(legacy_path.read_text(encoding="utf-8"))
                    if isinstance(payload, dict):
                        filename = str(payload.get("filename") or "documento_legacy").strip() or "documento_legacy"
                        legacy_chunks = payload.get("chunks") or []
                        if isinstance(legacy_chunks, list):
                            for i, text in enumerate(legacy_chunks):
                                chunks.append(
                                    {
                                        "doc_id": "legacy_doc",
                                        "filename": filename,
                                        "chunk_idx": i,
                                        "text": str(text),
                                    }
                                )
                except Exception:
                    pass

        return chunks

    # ── FAISS (caché en memoria por proceso) ──────────────────────────────────

    _faiss_cache: dict[str, tuple[object, float]] = {}  # area_id -> (vectorstore, mtime)

    def _faiss_stamp_path(self, area_id: str) -> Path:
        return self.base_dir / area_id / self.INDEX_DIR / "_stamp"

    def _invalidate_faiss(self, area_id: str) -> None:
        self._faiss_cache.pop(area_id, None)
        stamp = self._faiss_stamp_path(area_id)
        stamp.write_text(str(time.time()), encoding="utf-8")

    def _get_faiss(self, area_id: str, embeddings):
        """Obtiene o construye el vectorstore FAISS para un área."""
        stamp_path = self._faiss_stamp_path(area_id)
        stamp = float(stamp_path.read_text()) if stamp_path.exists() else 0.0
        cached = self._faiss_cache.get(area_id)
        if cached and cached[1] >= stamp:
            return cached[0]

        chunks = self._load_all_chunks(area_id)
        if not chunks:
            return None

        try:
            from langchain_community.vectorstores import FAISS as FAISSStore
            from langchain_core.documents import Document as LCDoc

            docs = [LCDoc(page_content=c["text"], metadata={"doc_id": c["doc_id"], "filename": c["filename"], "chunk_idx": c["chunk_idx"]}) for c in chunks]
            vs = FAISSStore.from_documents(docs, embeddings)
            self._faiss_cache[area_id] = (vs, time.time())
            return vs
        except Exception as e:
            print(f"⚠️ KnowledgeBase: Error construyendo FAISS para '{area_id}': {e}")
            return None

    # ── Búsqueda ─────────────────────────────────────────────────────────────

    def search(self, area_id: str, query: str, k: int = 5, embeddings=None) -> list[dict]:
        """
        Busca chunks relevantes en un área.
        Si hay embeddings usa FAISS, si no hace búsqueda por palabras clave.
        """
        chunks = self._load_all_chunks(area_id)
        if not chunks:
            return []

        if embeddings is not None:
            vs = self._get_faiss(area_id, embeddings)
            if vs is not None:
                try:
                    results = vs.similarity_search_with_score(query, k=k)
                    return [
                        {
                            "text": doc.page_content,
                            "filename": doc.metadata.get("filename", ""),
                            # FAISS devuelve distancia L2: menor = más similar.
                            # Negamos para que ordenar por mayor score = más relevante.
                            "score": -float(score),
                        }
                        for doc, score in results
                    ]
                except Exception:
                    pass

        # Fallback: búsqueda por palabras clave
        return self._keyword_search(chunks, query, k)

    def _keyword_search(self, chunks: list[dict], query: str, k: int) -> list[dict]:
        # Palabras vacías que no aportan discriminación
        stop_words = {"de", "la", "el", "en", "un", "una", "los", "las", "y", "a", "por", "para",
                      "con", "del", "al", "se", "que", "como", "es", "son", "su", "sus", "o",
                      "si", "no", "hay", "the", "and", "or", "is", "in", "to", "of", "a"}
        q_words = set(re.findall(r"\w+", query.lower())) - stop_words
        if not q_words:
            # Sin palabras discriminatorias → devolver primeros k chunks
            return [{"text": c["text"], "filename": c.get("filename", ""), "score": 0} for c in chunks[:k]]

        scored: list[tuple[float, dict]] = []
        for c in chunks:
            text_lower = c["text"].lower()
            hits = sum(1 for w in q_words if w in text_lower)
            scored.append((hits, c))

        scored.sort(key=lambda x: -x[0])

        # Tomar los k más relevantes; si hay empates con 0 hits aún incluirlos
        # (garantiza que siempre algún contexto llega al agente)
        return [{"text": c["text"], "filename": c.get("filename", ""), "score": sc} for sc, c in scored[:k]]

    def search_all_areas(self, query: str, k: int = 3, embeddings=None) -> list[dict]:
        """Busca en todas las áreas y devuelve los mejores resultados."""
        results: list[dict] = []
        for area in self.list_areas():
            area_id = area["id"]
            hits = self.search(area_id, query, k=k, embeddings=embeddings)
            for h in hits:
                h["area_name"] = area["name"]
                h["area_id"] = area_id
            results.extend(hits)
        results.sort(key=lambda x: x.get("score", 0), reverse=True)
        return results[:k * 2]

    def get_area_context(self, area_id: str, query: str, k: int = 4, embeddings=None) -> str:
        """Devuelve contexto formateado para inyectar en el prompt del agente."""
        hits = self.search(area_id, query, k=k, embeddings=embeddings)
        if not hits:
            return ""
        parts = [f"[Base de conocimiento - {self.get_area(area_id or '') or {}.get('name', area_id)}]"]
        for i, h in enumerate(hits, 1):
            parts.append(f"\n--- Fragmento {i} (fuente: {h['filename']}) ---\n{h['text']}")
        return "\n".join(parts)

    def get_full_context_for_query(self, query: str, k: int = 4, embeddings=None) -> str:
        """Busca en todas las áreas y devuelve contexto para el prompt."""
        hits = self.search_all_areas(query, k=k, embeddings=embeddings)
        if not hits:
            return ""
        parts = ["[Base de conocimiento - Procedimientos de empresa]"]
        for i, h in enumerate(hits, 1):
            area_label = h.get("area_name", h.get("area_id", ""))
            parts.append(f"\n--- Fragmento {i} | Área: {area_label} | Fuente: {h['filename']} ---\n{h['text']}")
        return "\n".join(parts)


# ── Singleton global ──────────────────────────────────────────────────────────
_manager: Optional[KnowledgeBaseManager] = None


def get_kb_manager() -> KnowledgeBaseManager:
    global _manager
    if _manager is None:
        _manager = KnowledgeBaseManager()
    return _manager
